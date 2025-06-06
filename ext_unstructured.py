from unstructured.partition.pdf import partition_pdf
import os
import sys
import pytesseract
from pdf2image import convert_from_path
import shutil
from pathlib import Path
import glob
import pandas as pd
from docx import Document
from openpyxl import load_workbook
import zipfile
from PIL import Image
import csv
import tempfile
import xlrd  # For old .xls files
import win32com.client  # For old .doc files
import concurrent.futures
import multiprocessing
from functools import partial
import json
import re  # Added for re module
import logging
import time
import gc
from datetime import datetime
from langdetect import detect
from collections import defaultdict

# Configure environment variables for Poppler and Tesseract
poppler_path = r"C:\Program Files\poppler-24.08.0\Library\bin"
tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
tessdata_path = r'C:\Program Files\Tesseract-OCR\tessdata'

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    'pdf': '.pdf',
    'word': ['.docx', '.doc'],
    'excel': ['.xlsx', '.xls'],
    'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'],
    'text': ['.txt', '.csv'],
    'compressed': ['.zip']
}

# Direct Tesseract configuration
os.environ['TESSDATA_PREFIX'] = tessdata_path
pytesseract.pytesseract.tesseract_cmd = tesseract_path

# Performance configurations
MAX_WORKERS = max(1, multiprocessing.cpu_count() - 1)  # Leave one CPU free
BATCH_SIZE = 4  # Number of pages to process in parallel for PDFs

def process_image(image):
    """Process a single image for OCR text extraction."""
    try:
        # Optimize image for OCR
        # Convert to grayscale
        if image.mode != 'L':
            image = image.convert('L')
        
        # Optional: Add image preprocessing here if needed
        # Example: image = image.point(lambda x: 0 if x < 128 else 255, '1')
        
        return pytesseract.image_to_string(image, lang='por')
    except Exception as e:
        print(f"Error in OCR processing: {str(e)}")
        return ""

def process_image_batch(images):
    """Process a batch of images in parallel."""
    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        results = list(executor.map(process_image, images))
    return results

def extract_from_doc(doc_path, output_txt_path):
    """Extract text from old .doc format using Word COM object."""
    try:
        # Convert .doc to .docx
        word = win32com.client.Dispatch("Word.Application")
        word.visible = False
        
        # Create a temporary file for docx
        temp_docx = os.path.join(tempfile.gettempdir(), "temp.docx")
        
        try:
            doc = word.Documents.Open(doc_path)
            doc.SaveAs2(temp_docx, FileFormat=16)  # 16 = docx format
            doc.Close()
            
            # Now process the docx file
            result = extract_from_docx(temp_docx, output_txt_path)
            
        finally:
            word.Quit()
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
                
        return result
    except Exception as e:
        print(f"Error processing DOC file {doc_path}: {str(e)}")
        return False

def extract_from_xls(xls_path, output_txt_path):
    """Extract text from old .xls format with clear table formatting."""
    try:
        workbook = xlrd.open_workbook(xls_path)
        text = []
        
        for sheet in workbook.sheets():
            text.append(f"\n{'='*80}")
            text.append(f"Sheet: {sheet.name}")
            text.append('='*80 + '\n')
            
            # Get maximum column widths
            col_widths = []
            for col in range(sheet.ncols):
                col_width = max(len(str(sheet.cell_value(row, col))) for row in range(sheet.nrows))
                col_widths.append(max(col_width, 3))  # minimum width of 3
            
            # Create table border
            separator = '+'
            for width in col_widths:
                separator += '-' * (width + 2) + '+'
            text.append(separator)
            
            # Process rows
            for row in range(sheet.nrows):
                row_text = []
                for col in range(sheet.ncols):
                    value = str(sheet.cell_value(row, col) if sheet.cell_value(row, col) != '' else '')
                    formatted_value = value.ljust(col_widths[col])
                    row_text.append(formatted_value)
                
                text.append('| ' + ' | '.join(row_text) + ' |')
                text.append(separator)
            
            text.append('')  # Add space between sheets
        
        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text))
        return True
    except Exception as e:
        print(f"Error processing XLS file {xls_path}: {str(e)}")
        return False

def extract_from_pdf(pdf_path, output_txt_path):
    """Extract text from a PDF file with clear table formatting."""
    try:
        # First try to extract tables using tabula-py
        import tabula
        
        print(f"\nExtracting tables from PDF: {pdf_path}")
        
        # Extract tables
        tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
        all_text = []
        
        # Process extracted tables
        if tables:
            print(f"Found {len(tables)} tables in PDF")
            for idx, table in enumerate(tables, 1):
                all_text.append(f"\n{'='*80}")
                all_text.append(f"Table {idx}")
                all_text.append('='*80 + '\n')
                
                # Convert DataFrame to string with proper formatting
                table_str = table.to_string(index=False, justify='left')
                # Add borders to the table
                lines = table_str.split('\n')
                width = max(len(line) for line in lines)
                border = '+' + '-'*(width-2) + '+'
                
                formatted_lines = []
                formatted_lines.append(border)
                for line in lines:
                    formatted_lines.append('| ' + line.ljust(width-4) + ' |')
                    formatted_lines.append(border)
                
                all_text.extend(formatted_lines)
                all_text.append('')  # Add space after table
        
        # Now extract regular text with OCR
        print("\nExtracting text with OCR...")
        images = convert_from_path(
            pdf_path,
            poppler_path=poppler_path,
            dpi=300
        )
        
        # Process images in batches
        for i in range(0, len(images), BATCH_SIZE):
            batch = images[i:i + BATCH_SIZE]
            print(f"Processing pages {i+1} to {min(i+BATCH_SIZE, len(images))}...")
            
            texts = process_image_batch(batch)
            for j, text in enumerate(texts, i+1):
                all_text.append(f"\n{'='*80}")
                all_text.append(f"Page {j}")
                all_text.append('='*80)
                all_text.append(text)
                all_text.append('')  # Add space between pages

        # Save the extracted text
        with open(output_txt_path, "w", encoding="utf-8") as f:
            f.write("\n".join(all_text))
        
        return True
    except Exception as e:
        print(f"Error processing PDF {pdf_path}: {str(e)}")
        # Fallback to basic OCR if table extraction fails
        return extract_from_pdf_basic(pdf_path, output_txt_path)

def extract_from_docx(docx_path, output_txt_path):
    """Extract text from a Word document."""
    try:
        doc = Document(docx_path)
        text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text.append(" | ".join(row_text))
        
        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text))
        return True
    except Exception as e:
        print(f"Error processing Word document {docx_path}: {str(e)}")
        return False

def extract_from_excel(excel_path, output_txt_path):
    """Extract text from an Excel file with clear table formatting."""
    try:
        wb = load_workbook(excel_path, data_only=True)
        text = []
        
        for sheet in wb:
            text.append(f"\n{'='*80}")
            text.append(f"Sheet: {sheet.title}")
            text.append('='*80 + '\n')
            
            # Get maximum column widths for formatting
            col_widths = []
            for col in sheet.columns:
                col_width = max(len(str(cell.value if cell.value is not None else '')) for cell in col)
                col_widths.append(max(col_width, 3))  # minimum width of 3
            
            # Process merged cells first
            merged_ranges = list(sheet.merged_cells.ranges)
            merged_cells = {}
            for cell_range in merged_ranges:
                value = sheet.cell(cell_range.min_row, cell_range.min_col).value
                for row in range(cell_range.min_row, cell_range.max_row + 1):
                    for col in range(cell_range.min_col, cell_range.max_col + 1):
                        merged_cells[(row, col)] = str(value if value is not None else '')

            # Create table border
            separator = '+'
            for width in col_widths:
                separator += '-' * (width + 2) + '+'
            text.append(separator)
            
            # Process rows
            for row in sheet.iter_rows():
                row_text = []
                for idx, cell in enumerate(row):
                    # Check if cell is part of a merged range
                    if (cell.row, cell.column) in merged_cells:
                        value = merged_cells[(cell.row, cell.column)]
                    else:
                        value = str(cell.value if cell.value is not None else '')
                    
                    # Format value to match column width
                    formatted_value = value.ljust(col_widths[idx])
                    row_text.append(formatted_value)
                
                text.append('| ' + ' | '.join(row_text) + ' |')
                text.append(separator)
            
            text.append('')  # Add space between sheets
        
        # Save formatted text
        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text))
            
        return True
    except Exception as e:
        print(f"Error processing Excel file {excel_path}: {str(e)}")
        return False

def extract_from_csv(csv_path, output_txt_path):
    """Extract text from a CSV file."""
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            text = []
            for row in reader:
                text.append(" | ".join(row))
        
        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text))
        return True
    except Exception as e:
        print(f"Error processing CSV file {csv_path}: {str(e)}")
        return False

def process_zip_file(zip_path, output_folder):
    """Extract and process files from a ZIP archive."""
    logger = logging.getLogger('text_extraction')
    logger.info(f"\nProcessing ZIP file: {zip_path}")
    
    try:
        # Validate ZIP file first
        if not os.path.exists(zip_path):
            logger.error(f"ZIP file does not exist: {zip_path}")
            return False
            
        if not zipfile.is_zipfile(zip_path):
            logger.error(f"Invalid ZIP file: {zip_path}")
            return False
            
        # Test ZIP file integrity
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            if zip_ref.testzip() is not None:
                logger.error(f"ZIP file is corrupted: {zip_path}")
                return False
        
        # Create temporary directory for extraction
        with tempfile.TemporaryDirectory() as temp_dir:
            logger.info(f"Created temporary directory for ZIP extraction")
            
            # Extract ZIP contents
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                # Get list of files before extraction
                file_list = zip_ref.namelist()
                if not file_list:
                    logger.warning(f"ZIP file is empty: {zip_path}")
                    return False
                    
                # Check for unsafe paths
                for file_path in file_list:
                    if file_path.startswith('/') or '..' in file_path:
                        logger.error(f"Unsafe path in ZIP: {file_path}")
                        return False
                
                # Safe to extract
                zip_ref.extractall(temp_dir)
                logger.info(f"Extracted {len(file_list)} files to temporary directory")
            
            # Process the extracted contents
            return process_folder(temp_dir, output_folder)
            
    except zipfile.BadZipFile as e:
        logger.error(f"Bad ZIP file {zip_path}: {str(e)}")
        return False
    except Exception as e:
        logger.error(f"Error processing ZIP file {zip_path}: {str(e)}")
        return False

def find_all_files(folder_path):
    """Find all supported files including those in ZIP files."""
    file_counts = {
        'pdf': 0,
        'word': 0,
        'excel': 0,
        'image': 0,
        'text': 0,
        'zip': 0,
        'total': 0
    }
    all_files = []
    zip_files = []
    
    print(f"\nScanning directory: {folder_path}")
    
    try:
        # First, find all files
        for root, dirs, files in os.walk(folder_path):
            print(f"\nScanning {root}")
            print(f"Found {len(files)} files in this directory")
            
            for file in files:
                try:
                    file_path = os.path.join(root, file)
                    file_ext = os.path.splitext(file)[1].lower()
                    
                    # Verify file exists and is readable
                    if not os.path.exists(file_path):
                        print(f"File does not exist: {file_path}")
                        continue
                        
                    if not os.access(file_path, os.R_OK):
                        print(f"File is not readable: {file_path}")
                        continue
                    
                    # Check if it's a ZIP file
                    if file_ext == '.zip':
                        print(f"Found ZIP file: {file}")
                        try:
                            with zipfile.ZipFile(file_path, 'r') as zf:
                                # Test ZIP file validity
                                if zf.testzip() is None:
                                    zip_files.append(file_path)
                                    file_counts['zip'] += 1
                                    # Count files inside ZIP
                                    zip_contents = zf.namelist()
                                    print(f"ZIP contains {len(zip_contents)} files")
                                    for zip_file in zip_contents:
                                        zip_ext = os.path.splitext(zip_file)[1].lower()
                                        if get_file_type(zip_file):
                                            file_type = get_file_type(zip_file)
                                            file_counts[file_type] += 1
                                            file_counts['total'] += 1
                                else:
                                    print(f"Invalid ZIP file: {file}")
                        except zipfile.BadZipFile:
                            print(f"Corrupt ZIP file: {file}")
                            continue
                    else:
                        # Check if it's a supported file type
                        file_type = get_file_type(file_path)
                        if file_type:
                            print(f"Found supported file: {file} (Type: {file_type})")
                            all_files.append(file_path)
                            file_counts[file_type] += 1
                            file_counts['total'] += 1
                        else:
                            print(f"Unsupported file type: {file}")
                except Exception as e:
                    print(f"Error processing file {file}: {str(e)}")
                    continue
        
        print(f"\n{'='*50}")
        print("File Count Summary:")
        print('='*50)
        print(f"PDFs: {file_counts['pdf']}")
        print(f"Word Documents: {file_counts['word']}")
        print(f"Excel Files: {file_counts['excel']}")
        print(f"Images: {file_counts['image']}")
        print(f"Text Files: {file_counts['text']}")
        print(f"ZIP Archives: {file_counts['zip']}")
        print(f"Total Processable Files: {file_counts['total']}")
        print('='*50)
        
        return all_files, zip_files, file_counts
    except Exception as e:
        print(f"Error scanning directory {folder_path}: {str(e)}")
        return [], [], file_counts

def get_file_type(file_path):
    """Determine the type of file based on its extension."""
    ext = os.path.splitext(file_path)[1].lower()
    
    for file_type, extensions in SUPPORTED_EXTENSIONS.items():
        if isinstance(extensions, list):
            if ext in extensions:
                return file_type
        elif ext == extensions:
            return file_type
    
    return None

def extract_text(file_path, output_txt_path):
    """Extract text from a file based on its type."""
    file_type = get_file_type(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    if file_type == 'pdf':
        return extract_from_pdf(file_path, output_txt_path)
    elif file_type == 'word':
        if ext == '.doc':
            return extract_from_doc(file_path, output_txt_path)
        else:
            return extract_from_docx(file_path, output_txt_path)
    elif file_type == 'excel':
        if ext == '.xls':
            return extract_from_xls(file_path, output_txt_path)
        else:
            return extract_from_excel(file_path, output_txt_path)
    elif file_type == 'image':
        return extract_from_image(file_path, output_txt_path)
    elif file_type == 'text':
        if file_path.lower().endswith('.csv'):
            return extract_from_csv(file_path, output_txt_path)
        else:
            shutil.copy2(file_path, output_txt_path)
            return True
    elif file_type == 'compressed':
        return process_zip_file(file_path, os.path.dirname(output_txt_path))
    else:
        print(f"Unsupported file type: {file_path}")
        return False

def process_folder(input_folder, output_folder):
    """Process all supported files in a folder and its subfolders."""
    logger = logging.getLogger('text_extraction')
    logger.info(f"\nProcessing folder: {input_folder}")
    
    try:
        # Validate input folder
        if not os.path.exists(input_folder):
            logger.error(f"Input folder does not exist: {input_folder}")
            return False
            
        if not os.path.isdir(input_folder):
            logger.error(f"Input path is not a directory: {input_folder}")
            return False
        
        # Create output folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        # Find all files recursively
        all_files = []
        empty_folders = []
        
        for root, dirs, files in os.walk(input_folder):
            has_valid_files = False
            
            for file in files:
                try:
                    file_path = os.path.join(root, file)
                    
                    # Basic file validation
                    if not os.path.exists(file_path):
                        logger.warning(f"File does not exist: {file_path}")
                        continue
                        
                    if not os.access(file_path, os.R_OK):
                        logger.warning(f"File is not readable: {file_path}")
                        continue
                    
                    # Handle nested ZIP files
                    if file.lower().endswith('.zip'):
                        logger.info(f"Found nested ZIP file: {file}")
                        if zipfile.is_zipfile(file_path):
                            # Get the parent folder name for the output structure
                            parent_folder = os.path.basename(os.path.dirname(file_path))
                            # Create output path that maintains the folder structure
                            zip_output = os.path.join(output_folder, parent_folder, os.path.splitext(file)[0])
                            os.makedirs(zip_output, exist_ok=True)
                            if process_zip_file(file_path, zip_output):
                                has_valid_files = True
                        else:
                            logger.warning(f"Invalid nested ZIP file: {file_path}")
                    else:
                        # Handle regular files
                        if get_file_type(file_path):
                            all_files.append(file_path)
                            has_valid_files = True
                
                except Exception as e:
                    logger.error(f"Error checking file {file}: {str(e)}")
                    continue
            
            if not has_valid_files and not any(dirs):
                empty_folders.append(root)
        
        # Process regular files
        successful = 0
        failed = 0
        
        for file_path in all_files:
            try:
                # Get the parent folder name for the output structure
                parent_folder = os.path.basename(os.path.dirname(file_path))
                # Create output path that maintains the folder structure
                output_folder_path = os.path.join(output_folder, parent_folder)
                os.makedirs(output_folder_path, exist_ok=True)
                
                # Create output file path
                file_name = os.path.splitext(os.path.basename(file_path))[0]
                output_path = os.path.join(output_folder_path, f"{file_name}.txt")
                
                logger.info(f"\nProcessing: {file_path}")
                logger.info(f"Output to: {output_path}")
                
                # Process with retry and check result
                result = process_with_retry(file_path, output_path)
                if result['success']:
                    if result['file_size'] > 0:
                        successful += 1
                        logger.info(f"Successfully processed: {file_path}")
                    else:
                        failed += 1
                        logger.warning(f"Zero-byte output for: {file_path}")
                else:
                    failed += 1
                    logger.error(f"Failed to process: {file_path}")
            
            except Exception as e:
                failed += 1
                logger.error(f"Error processing {file_path}: {str(e)}")
        
        # Print summary
        logger.info(f"\n{'='*80}")
        logger.info("Processing Complete!")
        logger.info(f"Total files processed: {len(all_files)}")
        logger.info(f"Successful: {successful}")
        logger.info(f"Failed: {failed}")
        logger.info(f"Empty folders found: {len(empty_folders)}")
        logger.info('='*80)
        
        return successful > 0 or len(empty_folders) > 0
        
    except Exception as e:
        logger.error(f"Error processing folder {input_folder}: {str(e)}")
        return False

def process_input(input_path, base_output_folder):
    """Process either a folder or ZIP file from the input path."""
    try:
        # Get the name to use for the output folder
        input_name = os.path.splitext(os.path.basename(input_path))[0]
        output_folder = os.path.join(base_output_folder, input_name)
        os.makedirs(output_folder, exist_ok=True)

        print(f"\nProcessing input: {input_path}")
        print(f"Output folder: {output_folder}")

        # Check if input is a ZIP file or directory
        if input_path.lower().endswith('.zip'):
            print(f"Processing as ZIP file...")
            return process_zip_file(input_path, output_folder)
        elif os.path.isdir(input_path):
            print(f"Processing as directory...")
            return process_folder(input_path, output_folder)
        else:
            print(f"Error: Input path {input_path} is neither a ZIP file nor a directory")
            return False
    except Exception as e:
        print(f"Error processing input {input_path}: {str(e)}")
        return False

def setup_logging():
    """Setup detailed logging configuration."""
    log_file = 'extraction.log'
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    return logging.getLogger('text_extraction')

def create_progress_tracker():
    """Create or load progress tracking system."""
    try:
        progress_file = 'extraction_progress.json'
        if os.path.exists(progress_file):
            with open(progress_file, 'r', encoding='utf-8') as f:
                progress = json.load(f)
                print("Loaded existing progress tracker")
                return progress
        
        # Initialize new progress tracker
        progress = {
            'total_files': 0,
            'processed_files': [],
            'failed_files': [],
            'remaining_files': [],
            'empty_folders': [],
            'last_processed_time': None,
            'last_processed_item': None,  # Track last processed item for resume
            'extraction_stats': {
                'successful': 0,
                'failed': 0,
                'total_time': 0,
                'avg_time_per_file': 0,
                'remaining_count': 0,  # Track remaining files
                'completion_percentage': 0  # Track completion percentage
            },
            'quality_metrics': {
                'avg_confidence': 0,
                'files_with_warnings': []
            }
        }
        return progress
    except Exception as e:
        print(f"Error creating/loading progress tracker: {str(e)}")
        return None

def update_progress_stats(progress, total_files):
    """Update progress statistics."""
    stats = progress['extraction_stats']
    stats['successful'] = len(progress['processed_files'])
    stats['failed'] = len(progress['failed_files'])
    stats['remaining_count'] = total_files - (stats['successful'] + stats['failed'])
    stats['completion_percentage'] = (stats['successful'] + stats['failed']) / total_files * 100 if total_files > 0 else 0
    return stats

def scan_for_files(input_path):
    """Scan input path and return all files to be processed."""
    logger = logging.getLogger('text_extraction')
    all_files = []
    
    try:
        for root, _, files in os.walk(input_path):
            for file in files:
                file_path = os.path.join(root, file)
                if file.lower().endswith('.zip'):
                    if zipfile.is_zipfile(file_path):
                        all_files.append(file_path)
                elif get_file_type(file_path):
                    all_files.append(file_path)
        
        logger.info(f"Found {len(all_files)} files to process")
        return all_files
    except Exception as e:
        logger.error(f"Error scanning for files: {str(e)}")
        return []

def save_progress(progress):
    """Save current progress to JSON file."""
    try:
        with open('extraction_progress.json', 'w', encoding='utf-8') as f:
            json.dump(progress, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logger = logging.getLogger('text_extraction')
        logger.error(f"Error saving progress: {str(e)}")

def analyze_text_accuracy(text):
    """Analyze text accuracy using various heuristics."""
    metrics = {
        'readability_score': 1.0,
        'structure_score': 1.0,
        'content_score': 1.0,
        'issues': []
    }
    
    try:
        # Check for common OCR errors
        ocr_errors = {
            'broken_words': len(re.findall(r'\w-\n\w', text)),  # Words broken across lines
            'misplaced_spaces': len(re.findall(r'\w{2,}\s{2,}\w{2,}', text)),  # Extra spaces between words
            'special_char_ratio': len(re.findall(r'[^a-zA-Z0-9\s.,!?;:()\-\'\"€$%]', text)) / (len(text) + 1),
            'number_errors': len(re.findall(r'\d[a-zA-Z]\d|\d\s\d(?!\d)', text))  # Numbers with letters in between
        }
        
        # Analyze text structure
        structure_issues = {
            'long_lines': len([line for line in text.split('\n') if len(line) > 100]),
            'empty_lines': text.count('\n\n'),
            'misaligned_text': len(re.findall(r'^\s+[A-Z]', text, re.MULTILINE))
        }
        
        # Check content quality
        content_issues = {
            'repeated_chars': len(re.findall(r'(.)\1{3,}', text)),  # More than 3 repeated characters
            'garbage_strings': len(re.findall(r'[^\s\n]{30,}', text)),  # Very long strings without spaces
            'low_letter_ratio': sum(c.isalpha() for c in text) / (len(text) + 1)
        }
        
        # Calculate scores
        # Readability score based on OCR errors
        metrics['readability_score'] = max(0, 1.0 - (
            0.1 * ocr_errors['broken_words'] +
            0.05 * ocr_errors['misplaced_spaces'] +
            0.3 * ocr_errors['special_char_ratio'] +
            0.1 * ocr_errors['number_errors']
        ))
        
        # Structure score based on layout issues
        metrics['structure_score'] = max(0, 1.0 - (
            0.01 * structure_issues['long_lines'] +
            0.02 * structure_issues['empty_lines'] +
            0.05 * structure_issues['misaligned_text']
        ))
        
        # Content score based on text quality
        metrics['content_score'] = max(0, 1.0 - (
            0.1 * content_issues['repeated_chars'] +
            0.2 * content_issues['garbage_strings'] +
            (1.0 - content_issues['low_letter_ratio'])
        ))
        
        # Add specific issues to the report
        if ocr_errors['broken_words'] > 0:
            metrics['issues'].append(f"Found {ocr_errors['broken_words']} broken words")
        if ocr_errors['special_char_ratio'] > 0.1:
            metrics['issues'].append("High ratio of special characters detected")
        if content_issues['garbage_strings'] > 0:
            metrics['issues'].append(f"Found {content_issues['garbage_strings']} potential garbage text segments")
        if content_issues['low_letter_ratio'] < 0.4:
            metrics['issues'].append("Low ratio of readable text detected")
            
    except Exception as e:
        metrics['issues'].append(f"Error analyzing text: {str(e)}")
        metrics['readability_score'] = 0.0
        metrics['structure_score'] = 0.0
        metrics['content_score'] = 0.0
    
    return metrics

def assess_text_quality(extracted_text, output_path=None):
    """Assess the quality of extracted text and check file size."""
    metrics = {
        'confidence_score': 1.0,  # Overall confidence score
        'warnings': [],
        'character_count': len(extracted_text),
        'word_count': len(extracted_text.split()),
        'potential_issues': [],
        'file_size': 0,
        'accuracy_metrics': None  # New field for accuracy metrics
    }
    
    # Check file size if path is provided
    if output_path and os.path.exists(output_path):
        file_size = os.path.getsize(output_path)
        metrics['file_size'] = file_size
        if file_size == 0:
            metrics['warnings'].append("Zero-byte file detected")
            metrics['confidence_score'] = 0
            return metrics
    
    # Check for empty text
    if len(extracted_text.strip()) == 0:
        metrics['warnings'].append("Empty text extracted")
        metrics['confidence_score'] = 0
        return metrics
    
    # Get detailed accuracy metrics
    accuracy_metrics = analyze_text_accuracy(extracted_text)
    metrics['accuracy_metrics'] = accuracy_metrics
    
    # Adjust confidence score based on accuracy metrics
    accuracy_score = (
        accuracy_metrics['readability_score'] * 0.4 +
        accuracy_metrics['structure_score'] * 0.3 +
        accuracy_metrics['content_score'] * 0.3
    )
    
    # Add accuracy issues to warnings
    metrics['warnings'].extend(accuracy_metrics['issues'])
    
    # Check for common OCR issues
    if '|' in extracted_text:
        metrics['warnings'].append("Possible table parsing issues")
        metrics['confidence_score'] -= 0.1
    
    # Check for garbled text
    non_ascii = sum(1 for c in extracted_text if ord(c) > 127)
    if non_ascii > len(extracted_text) * 0.1:  # More than 10% non-ASCII
        metrics['warnings'].append("High proportion of non-ASCII characters")
        metrics['confidence_score'] -= 0.2
    
    # Check word statistics
    words = extracted_text.split()
    if words:
        avg_word_length = sum(len(word) for word in words) / len(words)
        if avg_word_length < 2:
            metrics['warnings'].append("Words too short - possible OCR errors")
            metrics['confidence_score'] -= 0.3
        elif avg_word_length > 15:
            metrics['warnings'].append("Words too long - possible OCR errors")
            metrics['confidence_score'] -= 0.3
    
    # Try language detection
    try:
        lang = detect(extracted_text[:1000])  # Use first 1000 chars for speed
        if lang not in ['pt', 'en']:
            metrics['warnings'].append(f"Unexpected language detected: {lang}")
            metrics['confidence_score'] -= 0.2
    except:
        metrics['warnings'].append("Could not detect language")
        metrics['confidence_score'] -= 0.1
    
    # Combine original confidence score with accuracy score
    metrics['confidence_score'] = (metrics['confidence_score'] + accuracy_score) / 2
    
    # Ensure confidence score stays between 0 and 1
    metrics['confidence_score'] = max(0, min(1, metrics['confidence_score']))
    
    return metrics

def generate_processing_report(progress):
    """Generate a detailed processing report including accuracy metrics."""
    logger = logging.getLogger('text_extraction')
    
    try:
        report_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        report_file = f'processing_report_{report_time}.txt'
        
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write("TEXT EXTRACTION PROCESSING REPORT\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("="*80 + "\n\n")
            
            # Overall Statistics
            f.write("OVERALL STATISTICS\n")
            f.write("-"*50 + "\n")
            f.write(f"Total Files: {progress['total_files']}\n")
            f.write(f"Successfully Processed: {len(progress['processed_files'])}\n")
            f.write(f"Failed: {len(progress['failed_files'])}\n")
            f.write(f"Completion: {progress['extraction_stats']['completion_percentage']:.1f}%\n\n")
            
            # Accuracy Statistics
            f.write("ACCURACY METRICS\n")
            f.write("-"*50 + "\n")
            total_confidence = 0
            total_readability = 0
            total_structure = 0
            total_content = 0
            files_with_metrics = 0
            
            for file_path in progress['processed_files']:
                output_path = os.path.splitext(file_path)[0] + '.txt'
                if os.path.exists(output_path):
                    try:
                        with open(output_path, 'r', encoding='utf-8') as txt_file:
                            text = txt_file.read()
                            metrics = assess_text_quality(text, output_path)
                            
                            if metrics['accuracy_metrics']:
                                total_confidence += metrics['confidence_score']
                                total_readability += metrics['accuracy_metrics']['readability_score']
                                total_structure += metrics['accuracy_metrics']['structure_score']
                                total_content += metrics['accuracy_metrics']['content_score']
                                files_with_metrics += 1
                    except Exception as e:
                        logger.error(f"Error assessing file {file_path}: {str(e)}")
            
            if files_with_metrics > 0:
                f.write(f"Average Confidence Score: {total_confidence / files_with_metrics:.2%}\n")
                f.write(f"Average Readability Score: {total_readability / files_with_metrics:.2%}\n")
                f.write(f"Average Structure Score: {total_structure / files_with_metrics:.2%}\n")
                f.write(f"Average Content Score: {total_content / files_with_metrics:.2%}\n")
            
            # Failed Files
            if progress['failed_files']:
                f.write("\nFAILED FILES\n")
                f.write("-"*50 + "\n")
                for failed_file in progress['failed_files']:
                    f.write(f"- {failed_file}\n")
            
            f.write("\n" + "="*80 + "\n")
            f.write("End of Report\n")
            f.write("="*80 + "\n")
        
        logger.info(f"Processing report generated: {report_file}")
        
    except Exception as e:
        logger.error(f"Error generating processing report: {str(e)}")

def process_with_retry(file_path, output_path, max_retries=3, delay=5):
    """Process a file with retry mechanism and check for zero-byte outputs."""
    logger = logging.getLogger('text_extraction')
    
    for attempt in range(max_retries):
        try:
            start_time = time.time()
            success = extract_text(file_path, output_path)
            
            if success:
                # Check if file exists and has content
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    if file_size == 0:
                        logger.warning(f"Zero-byte file produced for {file_path}")
                        if attempt < max_retries - 1:
                            logger.info(f"Retrying... (Attempt {attempt + 2} of {max_retries})")
                            time.sleep(delay)
                            continue
                        else:
                            return {
                                'success': False,
                                'attempts': attempt + 1,
                                'error': "Zero-byte file produced",
                                'file_size': 0
                            }
                    
                    with open(output_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    quality_metrics = assess_text_quality(text, output_path)
                    
                    processing_time = time.time() - start_time
                    return {
                        'success': True,
                        'attempts': attempt + 1,
                        'processing_time': processing_time,
                        'quality_metrics': quality_metrics,
                        'file_size': file_size
                    }
                else:
                    raise Exception("Output file was not created")
            else:
                raise Exception("Extraction returned False")
                
        except Exception as e:
            logger.warning(f"Attempt {attempt + 1} failed for {file_path}: {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(delay)
    
    return {
        'success': False,
        'attempts': max_retries,
        'error': str(e),
        'file_size': 0 if os.path.exists(output_path) else None
    }

def enhanced_main():
    """Enhanced main processing function with resume capability."""
    logger = setup_logging()
    progress = create_progress_tracker()
    
    # Set the input path to the documents folder
    input_path = r"C:\Users\Miguel Carvalho\Documents\Github\dourogas-002\Ext_docs\documents"
    
    # Verify input path exists
    if not os.path.exists(input_path):
        logger.error(f"Input path does not exist: {input_path}")
        return
    
    # Create base output folder for all extractions
    base_output_folder = os.path.join(os.path.dirname(input_path), "unstructured_ext")
    os.makedirs(base_output_folder, exist_ok=True)
    
    logger.info(f"Processing documents folder: {input_path}")
    logger.info(f"Base output folder: {base_output_folder}")
    
    try:
        # Scan for all files first with detailed counting
        all_files, zip_files, file_counts = find_all_files(input_path)
        progress['total_files'] = file_counts['total']
        
        # Update remaining files list (exclude already processed and failed files)
        processed_set = set(progress['processed_files'])
        failed_set = set(progress['failed_files'])
        progress['remaining_files'] = [f for f in all_files if f not in processed_set and f not in failed_set]
        
        # Update statistics
        update_progress_stats(progress, file_counts['total'])
        
        # Log initial status with detailed breakdown
        logger.info("\nFile Processing Status:")
        logger.info("="*50)
        logger.info(f"Total files to process: {file_counts['total']}")
        logger.info(f"Already processed: {len(processed_set)}")
        logger.info(f"Previously failed: {len(failed_set)}")
        logger.info(f"Remaining to process: {len(progress['remaining_files'])}")
        logger.info("\nFile Type Breakdown:")
        logger.info(f"PDFs: {file_counts['pdf']}")
        logger.info(f"Word Documents: {file_counts['word']}")
        logger.info(f"Excel Files: {file_counts['excel']}")
        logger.info(f"Images: {file_counts['image']}")
        logger.info(f"Text Files: {file_counts['text']}")
        logger.info(f"ZIP Archives: {file_counts['zip']}")
        logger.info("="*50)
        
        # Process remaining files
        for item_path in progress['remaining_files']:
            try:
                # Update current item being processed
                progress['last_processed_item'] = item_path
                
                # Get the parent folder name for the output structure
                parent_folder = os.path.basename(os.path.dirname(item_path))
                # Create output folder path that maintains the folder structure
                output_folder = os.path.join(base_output_folder, parent_folder)
                os.makedirs(output_folder, exist_ok=True)
                
                logger.info(f"\nProcessing ({progress['extraction_stats']['completion_percentage']:.1f}% complete): {item_path}")
                
                if item_path.lower().endswith('.zip'):
                    # Handle ZIP files
                    with tempfile.TemporaryDirectory() as temp_dir:
                        logger.info(f"Extracting ZIP: {item_path}")
                        with zipfile.ZipFile(item_path, 'r') as zip_ref:
                            zip_ref.extractall(temp_dir)
                        if process_folder(temp_dir, output_folder):
                            progress['processed_files'].append(item_path)
                        else:
                            progress['failed_files'].append(item_path)
                
                elif os.path.isdir(item_path):
                    # Handle directories
                    if process_folder(item_path, output_folder):
                        progress['processed_files'].append(item_path)
                    else:
                        progress['failed_files'].append(item_path)
                
                else:
                    # Handle individual files
                    if get_file_type(item_path):
                        file_name = os.path.splitext(os.path.basename(item_path))[0]
                        output_path = os.path.join(output_folder, f"{file_name}.txt")
                        result = process_with_retry(item_path, output_path)
                        
                        if result['success']:
                            progress['processed_files'].append(item_path)
                        else:
                            progress['failed_files'].append(item_path)
                
                # Update and save progress after each item
                update_progress_stats(progress, file_counts['total'])
                progress['last_processed_time'] = datetime.now().isoformat()
                save_progress(progress)
                
                # Log progress
                logger.info(f"Progress: {progress['extraction_stats']['completion_percentage']:.1f}% complete")
                logger.info(f"Remaining: {progress['extraction_stats']['remaining_count']} files")
            
            except Exception as e:
                logger.error(f"Error processing {item_path}: {str(e)}")
                progress['failed_files'].append(item_path)
                save_progress(progress)
        
        logger.info("\nProcessing completed")
        update_progress_stats(progress, file_counts['total'])
        generate_processing_report(progress)
        
    except Exception as e:
        logger.error(f"Error during processing: {str(e)}")
    finally:
        # Save final progress and generate report
        progress['last_processed_time'] = datetime.now().isoformat()
        save_progress(progress)
        generate_processing_report(progress)

if __name__ == "__main__":
    enhanced_main()